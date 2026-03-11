import os
from docx import Document
from fpdf import FPDF
from .utils import (
    add_heading_paragraph,
    pdf_add_text,
    pdf_add_title,
    safe_filename
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "generated_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# -----------------------
# NDA Generator
# -----------------------
def generate_nda(call_llm, user, startup, other_party, purpose, preview=False):

    prompt = f"""
Generate a professional Non-Disclosure Agreement (NDA).

Company: {startup.startup_name}
Founder: {user.full_name}
Other Party: {other_party}
Purpose: {purpose}

Use clear headings and bullet-point clauses.
"""

    nda_text = call_llm(
        messages=[
            {"role": "system", "content": "You are a legal assistant."},
            {"role": "user", "content": prompt}
        ],
        model="mistral-small-latest"
    )

    file_name = f"NDA_{safe_filename(startup.startup_name)}_{safe_filename(other_party)}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)

    doc = Document()
    add_heading_paragraph(doc, "Non-Disclosure Agreement", nda_text)
    doc.save(file_path)

    if preview:
        return file_path, nda_text

    return file_path, nda_text


# -----------------------
# MoU Generator
# -----------------------
def generate_mou(call_llm, user, startup, partner_name, purpose, preview=False):

    prompt = f"""
Generate a Memorandum of Understanding (MoU).

Company: {startup.startup_name}
Founder: {user.full_name}
Partner: {partner_name}
Purpose: {purpose}

Use structured headings and bullet points.
"""

    mou_text = call_llm(
        messages=[
            {"role": "system", "content": "You are a legal assistant."},
            {"role": "user", "content": prompt}
        ],
        model="mistral-small-latest"
    )

    file_name = f"MoU_{safe_filename(startup.startup_name)}_{safe_filename(partner_name)}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)

    doc = Document()
    add_heading_paragraph(doc, "Memorandum of Understanding", mou_text)
    doc.save(file_path)

    if preview:
        return file_path, mou_text

    return file_path, mou_text


# -----------------------
# RTI Generator
# -----------------------
def generate_rti(call_llm, user, startup, authority, subject, purpose, preview=False):

    prompt = f"""
Draft a professional RTI application.

Applicant: {user.full_name}
Organization: {startup.startup_name}
Authority: {authority}
Subject: {subject}
Purpose: {purpose}

Use formal Indian RTI format.
"""

    rti_text = call_llm(
        messages=[
            {"role": "system", "content": "You are a legal assistant."},
            {"role": "user", "content": prompt}
        ],
        model="mistral-small-latest"
    )

    file_name = f"RTI_{safe_filename(startup.startup_name)}_{safe_filename(authority)}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)

    doc = Document()
    add_heading_paragraph(doc, "RTI Application", rti_text)
    doc.save(file_path)

    if preview:
        return file_path, rti_text

    return file_path, rti_text


# -----------------------
# Pitch Deck Generator (PDF)
# -----------------------
def generate_pitch_deck(call_llm, startup, preview=False):

    prompt = f"""
Create a concise startup pitch deck.

Startup: {startup.startup_name}

Sections:
• Problem
• Solution
• Market
• Business Model
• Team
• Vision

Use bullet points.
"""

    pitch_text = call_llm(
        messages=[
            {"role": "system", "content": "You are a startup assistant."},
            {"role": "user", "content": prompt}
        ],
        model="mistral-small-latest"
    )

    file_name = f"PitchDeck_{safe_filename(startup.startup_name)}.pdf"
    file_path = os.path.join(OUTPUT_DIR, file_name)

    pdf = FPDF()
    pdf.add_page()
    pdf_add_title(pdf, f"Pitch Deck – {startup.startup_name}")
    pdf_add_text(pdf, pitch_text)
    pdf.output(file_path)

    if preview:
        return file_path, pitch_text

    return file_path, pitch_text